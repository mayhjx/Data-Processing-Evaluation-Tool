from statistics import mean, stdev
from datetime import date

from docx import Document

from QCtool.models import BigDataFrame, PatientResult, QCframe, QCrecord
from QCtool.main import Evaluation


class VD(Evaluation):

    def __init__(self, f, instrument_num, instrument_type):
        self.f = f
        self.instrument_num = instrument_num
        self.instrument_type = instrument_type
        self.instrument_prefix = 'FXS-'
        self.D2 = "D2"
        self.D3 = "D3"
        self.D2_LLMI = 2.2
        self.D3_LLMI = 2.6
        self.separate = 'X' 
        if self.instrument_type == "AB":
            if self.f.name.endswith(".txt"):
                self.delimiter = '\t'
                self.Component = "Component Name" # 用于查找AB仪器数据目标化合物，Agilent和岛津无该变量
                self.NAME = 'Sample Name'
                self.ActualConc = "Actual Concentration"
                self.CONC = 'Calculated Concentration'
                self.ACCURACY = 'Accuracy'
                self.SN = 'Signal / Noise'
                self.USED = 'Used'
                self.R = 'Correlation Coefficient'
                self.Area = 'Area'
                self.ISArea = 'IS Area'
                self.IonRatio = None
                self.RT = 'Retention Time'
                self.ISRT = 'IS Retention Time'
                self.LOCATE = 'X'  # 定位孔实验号前缀
            elif self.f.name.endswith(".docx"):
                self.NAME = 'Sample Name'
                self.ActualConc = "Target  [Conc].(nmol/L)"
                self.CONC = 'Calculated Conc.(nmol/L)'
                self.ACCURACY = "Accuracy(%)"
                self.SN = 'S/N '
                self.USED = '[Use Record]'
                self.R = 'R2'
                self.Area = 'Area (cps)'
                self.ISArea = 'IS Area(cps)'
                self.RT = 'RT (min)'
                self.ISRT = 'IS Retention Time (min)'
                self.IonRatio = "MRM Ratio"
                self.LOCATE = 'X'  # 定位孔实验号前缀
        elif self.instrument_type == "Agilent":
            self.delimiter = ','
            self.NAME = 'Sample Name'
            self.ActualConc = None
            self.CONC = 'Calc. Conc.'
            self.ACCURACY = 'Accuracy'
            self.SN = 'S/N'
            self.USED = None
            self.R = 'CF R2'
            self.Area = 'Resp.'
            self.ISArea = 'IS Resp.'
            self.IonRatio = "Ion Ratio"
            self.RT = 'RT'
            self.ISRT = 'IS RT'
            self.LOCATE = 'X'  # 定位孔实验号前缀
        elif self.instrument_type == "daojin":
            self.delimiter = '\t'
            self.NAME = '样品名'
            self.ActualConc = "标准浓度"
            self.CONC = '浓度'
            self.ACCURACY = '精确度%'
            self.SN = 'S/N'
            self.USED = '校准点'  # true:'*'
            self.R = None
            self.Area = '面积'
            self.ISArea = '内标面积'
            self.IonRatio = None
            self.RT = '保留时间'
            self.ISRT = '内标保留时间'
            self.LOCATE = 'X'  # 定位孔实验号前缀

        super().__init__(self.NAME,
                         self.ActualConc,
                         self.CONC,
                         self.RT,
                         self.ISRT,
                         self.SN,
                         self.Area,
                         self.ISArea,
                         self.R,
                         self.USED,
                         self.ACCURACY,
                         self.IonRatio,
                         self.LOCATE,
                        )

    def reader(self, decoding='utf-8'):

        # InMemoryFile to list
        data = list()
        for line in self.f.read().splitlines():
            data.append(line.decode(decoding).split(self.delimiter))
        return data

    def AB_data_format(self):

        data = self.reader()
        try:
            head = data[0]
            conc = head.index(self.CONC)
            comp = head.index(self.Component)
        except:
            return

        dic = dict()
        for d in data[-1:0:-1]:
            if d[conc] in ["degenerate", "N/A", ]:
                data.remove(d)

        dic[self.D2] = list()
        dic[self.D2].append(head)
        dic[self.D2].extend([d for d in data if self.D2 in d[comp]])

        dic[self.D3] = list()
        dic[self.D3].append(head)
        dic[self.D3].extend([d for d in data if self.D3 in d[comp]])

        return dic

    def Agilent_data_format(self):

        data = self.reader()
        
        row1 = data[0]
        row2 = data[1]

        head = list()
        title = ''
        for h in list(zip(row1, row2)):
            if h[0]:
                title = h[0].split(' ')[0]
                # 将原始的标题行组合在一起，再替换为项目简称
                # 内标的Area加上IS与样品的Area区分
                if self.D2 in title:
                    if "IS" in title:
                        title = self.D2 + " IS"
                    else:
                        title = self.D2
                if "Qualifier (413.3 -> 355.2)" in h[0]:
                    title = self.D2 + " Ion"
                    
                elif self.D3 in title:
                    if "IS" in title:
                        title = self.D3 + " IS"
                    else:
                        title = self.D3
                if "Qualifier (401.3 -> 365.2)" in h[0]:
                    title = self.D3 + " Ion"
            h = "{} {}".format(title, h[1])
            head.append(h)

        data.remove(row1)
        data.remove(row2)
        data.insert(0, head)
        
        # *转换行列*
        transpose_data = list(map(list, zip(*data)))

        # 前3行是Sample信息
        transpose_D2 = transpose_data[0:2]
        transpose_D3 = transpose_data[0:2]

        for d in transpose_data:
            if self.D2 in d[0]:
                # 去掉之前加上的D2，下同
                d[0] = d[0][len(self.D2) + 1:]
                transpose_D2.append(d)
            elif self.D3 in d[0]:
                d[0] = d[0][len(self.D3) + 1:]
                transpose_D3.append(d)
        
        dic = dict()
        dic[self.D2] = list()
        dic[self.D2].extend(list(map(list, zip(*transpose_D2))))

        dic[self.D3] = list()
        dic[self.D3].extend(list(map(list, zip(*transpose_D3))))

        return dic

    def daojin_data(self):
        # 岛津数据转换
        data = self.reader('ansi')

        dic = dict()
        for d in data:
            if d[0] == 'Name' and 'D2' in d[1]:
                k = self.D2
                dic[k] = list()
            elif d[0] == 'Name' and 'D3' in d[1]:
                k = self.D3
                dic[k] = list()

            # 最后一个元素不为空则为样品数据
            if d[0] not in ['ID#', 'Name'] and d[-1]:
                dic[k].append(d)
        
        return dic

    def AB_data_format_docx(self):

        docx = Document(self.f)

        dic = dict()
        dic[self.D2] = list()
        dic[self.D3] = list()

        # 获取线性R2
        D2R2 = None
        D3R2 = None
        for paragraph in docx.paragraphs:
            if self.D2 in paragraph.text and not D2R2:
                D2R2 = paragraph.text.split("Correlation coefficient:")[1]
            elif self.D3 in paragraph.text and not D3R2:
                D3R2 = paragraph.text.split("Correlation coefficient:")[1]

        D2table = docx.tables[1]
        D3table = docx.tables[5]


        def read_table_data(table):
            # table 是一个Document.table对象
            # 优化读取table的速度，由8s优化到0.12s
            # if not isinstance(Document.tables, table): return
            cells = table._cells
            ROWS = len(table.rows)
            COLUMNS = len(table.columns)
            data = []
            datas = []
            for i in range(ROWS* COLUMNS):
                text = cells[i].text.replace("\n", "")
                if i % 12 != 0 or i == 0:
                    data.append(text)
                    continue
                else:
                    if data[0] == self.NAME:
                        data.append(self.R)
                    else:
                        data.append(D2R2)
                    datas.append(data)
                    data = []
                    data.append(text)
            datas.append(data)

            return datas

        dic[self.D2] = read_table_data(D2table)
        dic[self.D3] = read_table_data(D3table)

        # D2table = docx.tables[1]
        # for row in D2table.rows:
        #     data = []
        #     for cell in row.cells:
        #         text = cell.text.replace("\n", "")
        #         data.append(text)
        #     # 在末尾加R值
        #     if data[0] == self.NAME:
        #         data.append(self.R)
        #     else:
        #         data.append(D2R2)
        #     dic[self.D2].append(data)

        # D3table = docx.tables[5]
        # for row in D3table.rows:
        #     data = []
        #     for cell in row.cells:
        #         text = cell.text.replace("\n", "")
        #         data.append(text)
        #     if data[0] == self.NAME:
        #         data.append(self.R)
        #     else:
        #         data.append(D3R2)
        #     dic[self.D3].append(data)  

        return dic


    def get_dic_data(self):

        if self.instrument_type == "AB":
            if self.f.name.endswith(".txt"):
                dic = self.AB_data_format()
            elif self.f.name.endswith(".docx"):
                dic = self.AB_data_format_docx()
        elif self.instrument_type == "Agilent":
            dic = self.Agilent_data_format()
        elif self.instrument_type == "daojin":
            dic = self.daojin_data()

        return dic

    def get_LLMI(self, data):

        '''计算STD0峰面积与STD1峰面积一半的比值：
        1. 比值<10%: 10分
        2. 10%<=比值<15%: 9分
        3. 15%<=比值<20%: 8分
        4. 比值>=20%: 0分'''

        dic = dict()
        dic['score'] = 0
        dic['violation'] = list()

        samplename = self.index(data, self.SampleName)
        area = self.index(data, self.Area)

        if samplename is None or not area:
            dic['score'] = None
            return dic

        for std in data[1:]:
            if std[samplename].upper().startswith("STD") and std[samplename].upper().endswith("0"):
                std0_area = float(std[area])
                break

        for std in data[1:]:
            if std[samplename].upper().startswith("STD") and std[samplename].upper().endswith("1"):
                std1_area = float(std[area])
                break

        if std0_area is None or std1_area is None:
            dic['score'] = None
            return dic

        percent = round(std0_area / std1_area * 2, 2)

        if percent < 0.1:
            dic['score'] = 10
            if self.DEBUG: dic['violation'].append('比值: {:.0%}'.format(percent))
        elif percent < 0.15:
            dic['score'] = 9
            dic['violation'].append('比值: {:.0%}'.format(percent))
        elif percent < 0.2:
            dic['score'] = 8
            dic['violation'].append('比值: {:.0%}'.format(percent))
        elif percent >= 0.2:
            dic['score'] = 0
            dic['violation'].append('比值: {:.0%}'.format(percent))

        return dic


    def split_plate_data(self, data):

        # 传进结果列表data
        # 获取实验号位置
        # 初始化字典和板号（0）
        # 根据分隔实验号区分板,将样品结果放进字典（板号：结果列表）
        # 每次遇到该特殊实验号则板号加1
        # 返回字典（板号：结果列表）

        # 板号0存放曲线信息

        head = data[0]
        separate = self.separate
        samplename = data[0].index(self.NAME)

        plate = dict()
        p = 0  # key name
        plate[p] = list()
        plate[p].append(head)
        
        for d in data[1:]:
            if separate in d[samplename].upper():
                p = d[samplename].replace(separate,'')
                plate[p] = list()
                plate[p].append(head)
            plate[p].append(d)

        return plate

    def get_totalDmean(self, data):
        
        # data参数为字典,包含D2和D3数据
        # {data: {D2: [sample data],
        #             [sample data],
        #               ...},
        #       {D3: [sample data],
        #             [sample data],
        #               ...}
        # 计算D2加D3的均值(小于LLMI算为0)
        # 返回总D（实验号，总D值）

        samplename = data['D2'][0].index(self.NAME)
        conc = data['D2'][0].index(self.CONC)

        totalD = dict()

        # D2 D3小于LLMI 算为0
        for d in data[self.D2]:
            if not d[samplename].upper().startswith(("BLANK", "TEST", "STD", "QC",)):
                name = d[samplename]
                value = d[conc]
                try:
                    if value == '< 0' or float(value) < self.D2_LLMI:
                        totalD[name] = 0
                    else:
                        totalD[name] = float(value)
                except:
                    totalD[name] = value

        for d in data[self.D3][1:]:
            if not d[samplename].upper().startswith(("BLANK", "TEST", "STD", "QC",)):
                name = d[samplename]
                value = d[conc]
                if value == '< 0' or float(value) < self.D3_LLMI:
                    totalD[name] += 0
                else:
                    totalD[name] += float(value)

        totalDlist = [[name, value] for name, value in totalD.items()]

        return totalDlist

    def patient_result(self, data):

        # 初始化一个dict存放各板均值和违反点
        # 剔除大于mean±3.09SD的数据
        # 在大数据计算接受范围内（Mean±SD），以板为单位，计算总D的均值
        # 均值<=Mean±1SD: 10分
        # 均值<=Mean±2SD: 9分
        # 均值<=Mean±3SD: 8分
        # 均值>Mean±3SD: 0分
        '''以板为单位，剔除超出该板mean±3.09*sd的数据后计算总D的均值，
            与当月的大数据（Mean±SD）接受范围判断：
        1. 均值<=Mean±1SD: 10分
        2. Mean±1SD<均值<=Mean±2SD: 9分
        3. Mean±2SD<均值<=Mean±3SD: 8分
        4. 均值>Mean±3SD: 0分'''


        dic = dict()
        
        per_plate_data = self.split_plate_data(data)

        for plate, data in per_plate_data.items():

            if plate == 0: continue
            
            dic[plate] = dict()
            platedic = dic[plate]
            platedic['score'] = 0
            platedic['mean'] = 0
            platedic['violation'] = list()

            sampledata = [d[1] for d in data[1:] if not d[0].upper().startswith(("BLANK", "TEST", "STD", "QC", self.separate,))]
            ave = mean(sampledata)

            if len(sampledata) > 1:
                sd = stdev(sampledata)
            else:
                sd = 0

            # 剔除大于X ± 3.09SD的值
            for sample in sampledata:
                if abs(sample - ave) > 3.09 * sd:
                    sampledata.remove(sample)

            ave = round(mean(sampledata), 2)

            # 样品数量
            platedic['num'] = len(sampledata)
            
            platedic['mean'] = "{:.2f}".format(ave)

            # 相对大数据均值的偏倚(保留一位小数)
            big_data_X = self.get_big_data_mean()
            big_data_SD = self.get_big_data_sd()
            if big_data_SD == 0 or big_data_X == 0:
                continue

            rsd = round((ave - big_data_X) / big_data_SD, 1)

            if abs(rsd) <= 1:
                platedic['score'] = 10
                platedic['violation'].append("/")
            elif abs(rsd) <= 2:
                platedic['score'] = 9
                platedic['violation'].append(">1SD")
            elif abs(rsd) <= 3:
                platedic['score'] = 8
                platedic['violation'].append(">2SD")
            elif abs(rsd) > 3:
                platedic['score'] = 0
                platedic['violation'].append(">3SD")

            # rsd结果放进数据库
            PatientResult.objects.create(instrument=self.instrument_num, platenum=plate, num=len(sampledata), mean=ave, sd=rsd)

        return dic

    def get_big_data_mean(self):
        # 返回当前月份的大数据均值
        currentmonth = date.today().strftime("%m")
        currentframe = BigDataFrame.objects.filter(month=currentmonth)
        if currentframe.exists(): 
            mean = currentframe[0].mean
            return mean
        else:
            return

    def get_big_data_sd(self):
        # 返回当前月份的大数据标准差
        currentmonth = date.today().strftime("%m")
        currentframe = BigDataFrame.objects.filter(month=currentmonth)
        if currentframe.exists(): 
            sd = currentframe[0].sd
            return sd
        else:
            return
        
    def output(self):

        dic = self.get_dic_data()
        if not dic: 
            raise IOError("无法识别原始数据")
        
        output = dict()

        output["accuracy_doc"] = self.get_accuracy.__doc__
        output["r_doc"] = self.get_r.__doc__
        output["LLMI_doc"] = self.get_LLMI.__doc__
        output["sn_doc"] = self.get_sn.__doc__
        output["isarea_doc"] = self.get_isarea.__doc__
        output["rt_doc"] = self.get_rt.__doc__
        output["isrt_doc"] = self.get_isrt.__doc__
        output["qc_doc"] = self.get_qc.__doc__
        output["locate_doc"] = self.get_locate.__doc__
        output["patient_result_doc"] = self.patient_result.__doc__
        output["ion_ratio_doc"] = self.get_Ion_Ratio.__doc__

        # 病人结果分布打分
        data = self.get_totalDmean(dic)
        output["patient_result"] = list()
        output["patient_result"].append(self.patient_result(data))

        # 大数据框架
        # output['big_data_mean'] = self.get_big_data_mean()
        # output['big_data_sd'] = self.get_big_data_sd()

        # 按板号分数据
        D2_per_plate_data = self.split_plate_data(dic[self.D2])
        D3_per_plate_data = self.split_plate_data(dic[self.D3])

        output["d2_isarea"] = list()
        output["d2_rt"] = list()
        output["d2_isrt"] = list()
        output["d2_qc"] = list()
        output["d2_locate"] = list()
        output['d2_total_score'] = list()

        output["d3_isarea"] = list()
        output["d3_rt"] = list()
        output["d3_isrt"] = list()
        output["d3_qc"] = list()
        output["d3_locate"] = list()
        output['d3_total_score'] = list()

        for plate, data in D2_per_plate_data.items():

            if plate == 0:
                output['d2_accuracy'] = self.get_accuracy(data)
                output['d2_r'] = self.get_r(data)
                output["d2_LLMI"] = self.get_LLMI(data)
                output["d2_sn"] = self.get_sn(data)

                # 板0分数
                plate0_score = (output['d2_accuracy']['score'] or 0) + \
                            (output['d2_r']['score'] or 0) + \
                            (output['d2_LLMI']['score'] or 0) + \
                            (output['d2_sn']['score'] or 0)
            else:
                isarea = self.get_isarea(data)
                rt = self.get_rt(data)
                isrt = self.get_isrt(data)
                qc = self.get_qc(data, self.instrument_prefix + self.instrument_num, self.D2)
                locate = self.get_locate(data)

                output["d2_isarea"].append(dict([(plate, isarea)]))
                output["d2_rt"].append(dict([(plate, rt)]))
                output["d2_isrt"].append(dict([(plate, isrt)]))
                output["d2_qc"].append(dict([(plate, qc)]))
                output["d2_locate"].append(dict([(plate, locate)]))

                # 每板分数
                plate_score = (isarea['score'] or 0) + \
                             (rt['score'] or 0) + \
                             (isrt['score'] or 0) + \
                             (qc['score'] or 0) + \
                             (locate['score'] or 0) + \
                             (output['patient_result'][0][plate]['score'] or 0)

                output['d2_total_score'].append(dict([(plate, plate0_score+plate_score)]))
                
        for plate, data in D3_per_plate_data.items():

            if plate == 0:
                output['d3_accuracy'] = self.get_accuracy(data)
                output['d3_r'] = self.get_r(data)
                output["d3_LLMI"] = self.get_LLMI(data)
                output["d3_sn"] = self.get_sn(data)

                # 板0 分数
                plate0_score = (output['d3_accuracy']['score'] or 0) + \
                            (output['d3_r']['score'] or 0) + \
                            (output['d3_LLMI']['score'] or 0) + \
                            (output['d3_sn']['score'] or 0)

            else:
                isarea = self.get_isarea(data)
                rt = self.get_rt(data)
                isrt = self.get_isrt(data)
                qc = self.get_qc(data, self.instrument_prefix + self.instrument_num, self.D3)
                locate = self.get_locate(data)

                output["d3_isarea"].append(dict([(plate, isarea)]))
                output["d3_rt"].append(dict([(plate, rt)]))
                output["d3_isrt"].append(dict([(plate, isrt)]))
                output["d3_qc"].append(dict([(plate, qc)]))
                output["d3_locate"].append(dict([(plate, locate)]))
                
                # 每板分数
                plate_score = (isarea['score'] or 0) + \
                             (rt['score'] or 0) + \
                             (isrt['score'] or 0) + \
                             (qc['score'] or 0) + \
                             (locate['score'] or 0) + \
                             (output['patient_result'][0][plate]['score'] or 0)


                output['d3_total_score'].append(dict([(plate, plate0_score+plate_score)]))

        output["d2_Ion_Ratio"] = self.get_Ion_Ratio(dic[self.D2], self.D2_LLMI, self.instrument_prefix + self.instrument_num, self.D2)
        output["d3_Ion_Ratio"] = self.get_Ion_Ratio(dic[self.D3], self.D3_LLMI, self.instrument_prefix + self.instrument_num, self.D3)
        return output